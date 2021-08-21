from docx import Document
from docx.shared import Inches,Cm,Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT#导入的关于word操作的相关模块
import datetime #获取当前日期
import os

path = r'E:\文件资料\2021814处理微水震荡数据\微水实验报告.docx'#docx文件的地址
docx = Document(path)
tables = docx.tables#获取所有表格
table=tables[0]#获取第一个表格
filenames = os.listdir(r'E:\文件资料\2021814处理微水震荡数据\水位恢复图')
i = 1
for name in filenames:
    if i%2 != 0:
        位置 = table.cell(int((i-1)/2),0).paragraphs[0]
    elif i%2 == 0:
        位置 = table.cell(int(i/2-1),1).paragraphs[0]
    run =位置.add_run()#插入一个文字块
    photo_path = rf"E:\文件资料\2021814处理微水震荡数据\水位恢复图\{name}"
    picture =run.add_picture(photo_path)
    picture.height=Cm(4.88)#设置图片高度
    picture.width=Cm(7.29 )#设置图片宽度
    i += 1
    
docx.save(path)
